import os
import pytest
from httpx import AsyncClient
from app.domain.models.user import UserRole

@pytest.mark.asyncio
async def test_document_upload_and_delete_flow(client: AsyncClient):
    # 1. Register and login as Admin (restricted endpoint check)
    reg_response = await client.post(
        "/api/v1/auth/register",
        json={
            "email": "admin@example.com",
            "password": "adminsecurepwd",
            "full_name": "Admin User",
            "role": "ADMIN"
        }
    )
    assert reg_response.status_code == 201
    
    login_response = await client.post(
        "/api/v1/auth/token",
        data={
            "username": "admin@example.com",
            "password": "adminsecurepwd"
        }
    )
    assert login_response.status_code == 200
    token = login_response.json()["access_token"]
    headers = {"Authorization": f"Bearer {token}"}
    
    # 2. Register a standard Employee user to verify auth restrictions
    emp_reg = await client.post(
        "/api/v1/auth/register",
        json={
            "email": "employee@example.com",
            "password": "empsecurepwd",
            "full_name": "Employee User",
            "role": "EMPLOYEE"
        }
    )
    assert emp_reg.status_code == 201
    
    emp_login = await client.post(
        "/api/v1/auth/token",
        data={
            "username": "employee@example.com",
            "password": "empsecurepwd"
        }
    )
    assert emp_login.status_code == 200
    emp_token = emp_login.json()["access_token"]
    emp_headers = {"Authorization": f"Bearer {emp_token}"}

    # 3. Test that Employee CANNOT upload documents
    files = {"file": ("policy.txt", b"This is a corporate HR policy on leaves.", "text/plain")}
    emp_upload_response = await client.post(
        "/api/v1/documents/",
        files=files,
        headers=emp_headers
    )
    assert emp_upload_response.status_code == 403 # Forbidden for Employees
    
    # 4. Upload document as Admin
    files = {"file": ("policy.txt", b"This is a corporate HR policy on leaves.", "text/plain")}
    upload_response = await client.post(
        "/api/v1/documents/",
        files=files,
        headers=headers
    )
    assert upload_response.status_code == 202 # Accepted
    doc_data = upload_response.json()
    assert doc_data["filename"] == "policy.txt"
    assert doc_data["status"] == "PROCESSING"
    doc_id = doc_data["id"]
    
    # Wait for the background task to complete processing and embedding
    import asyncio
    completed = False
    for _ in range(10): # try up to 10 times (1 second total)
        details_response = await client.get(f"/api/v1/documents/{doc_id}", headers=headers)
        assert details_response.status_code == 200
        details_data = details_response.json()
        if details_data["status"] == "COMPLETED":
            completed = True
            break
        await asyncio.sleep(0.1)
    assert completed, "Background document ingestion timed out"

    
    # 5. List documents
    list_response = await client.get("/api/v1/documents/", headers=headers)
    assert list_response.status_code == 200
    list_data = list_response.json()
    assert len(list_data) >= 1
    assert any(d["id"] == doc_id for d in list_data)

    # 6. Delete document as Admin
    delete_response = await client.delete(f"/api/v1/documents/{doc_id}", headers=headers)
    assert delete_response.status_code == 204
    
    # Try fetching details again -> should return 404
    missing_response = await client.get(f"/api/v1/documents/{doc_id}", headers=headers)
    assert missing_response.status_code == 404


@pytest.mark.asyncio
async def test_docx_document_upload_and_parse(client: AsyncClient):
    # 1. Register and login as Admin
    reg_response = await client.post(
        "/api/v1/auth/register",
        json={
            "email": "admin_docx@example.com",
            "password": "adminsecurepwd",
            "full_name": "Admin Docx User",
            "role": "ADMIN"
        }
    )
    assert reg_response.status_code == 201
    
    login_response = await client.post(
        "/api/v1/auth/token",
        data={
            "username": "admin_docx@example.com",
            "password": "adminsecurepwd"
        }
    )
    assert login_response.status_code == 200
    token = login_response.json()["access_token"]
    headers = {"Authorization": f"Bearer {token}"}

    # 2. Create a mock word docx in bytes
    import io
    import docx
    doc = docx.Document()
    doc.add_paragraph("This is a corporate HR policy on travel reimbursement.")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Meal allowance"
    table.rows[0].cells[1].text = "$50 per day"
    
    file_bytes = io.BytesIO()
    doc.save(file_bytes)
    file_bytes.seek(0)
    
    files = {"file": ("travel_policy.docx", file_bytes.read(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    
    # 3. Upload Docx
    upload_response = await client.post(
        "/api/v1/documents/",
        files=files,
        headers=headers
    )
    assert upload_response.status_code == 202
    doc_id = upload_response.json()["id"]
    
    # 4. Wait for processing completion (completed state check)
    import asyncio
    completed = False
    for _ in range(10):
        details_response = await client.get(f"/api/v1/documents/{doc_id}", headers=headers)
        details_data = details_response.json()
        if details_data["status"] == "COMPLETED":
            completed = True
            break
        await asyncio.sleep(0.1)
    assert completed
    
    # 5. Clean up
    delete_response = await client.delete(f"/api/v1/documents/{doc_id}", headers=headers)
    assert delete_response.status_code == 204
