import os
import uuid
import logging
from typing import List, Optional, Tuple
from app.core.chunker import RecursiveCharacterTextSplitter
from app.core.exceptions import EntityNotFoundError
from app.domain.interfaces.repositories import IDocumentRepository, IVectorRepository
from app.domain.models.document import DocumentCreate, DocumentResponse, DocumentStatus
from app.services.embedding_service import EmbeddingService

logger = logging.getLogger(__name__)

class DocumentService:
    def __init__(
        self,
        doc_repo: IDocumentRepository,
        vector_repo: IVectorRepository,
        embedding_service: EmbeddingService
    ):
        self.doc_repo = doc_repo
        self.vector_repo = vector_repo
        self.embedding_service = embedding_service
        self.splitter = RecursiveCharacterTextSplitter(chunk_size=600, chunk_overlap=60)

    async def list_documents(self, user_id: Optional[uuid.UUID] = None) -> List[DocumentResponse]:
        return await self.doc_repo.list_all(user_id)

    async def get_document(self, doc_id: uuid.UUID) -> DocumentResponse:
        doc = await self.doc_repo.get_by_id(doc_id)
        if not doc:
            raise EntityNotFoundError(f"Document with ID {doc_id} not found")
        return doc

    async def delete_document(self, doc_id: uuid.UUID) -> bool:
        doc = await self.doc_repo.get_by_id(doc_id)
        if not doc:
            raise EntityNotFoundError(f"Document with ID {doc_id} not found")
            
        # 1. Delete from vector DB
        await self.vector_repo.delete_chunks(doc_id)
        
        # 2. Delete file from local disk storage
        if os.path.exists(doc.storage_path):
            try:
                os.remove(doc.storage_path)
            except Exception as e:
                logger.error(f"Failed to delete file from disk storage: {e}")
                
        # 3. Delete from SQL DB
        return await self.doc_repo.delete(doc_id)

    async def create_pending_document(
        self,
        filename: str,
        storage_path: str,
        content_type: str,
        file_size: int,
        uploaded_by: uuid.UUID
    ) -> DocumentResponse:
        doc_create = DocumentCreate(
            filename=filename,
            content_type=content_type,
            file_size=file_size,
            storage_path=storage_path,
            uploaded_by=uploaded_by
        )
        return await self.doc_repo.create(doc_create)

    async def parse_and_embed_document(
        self,
        doc_id: uuid.UUID,
        storage_path: str,
        content_type: str,
        filename: str
    ) -> None:
        try:
            # 1. Extract text and page numbers from file
            pages = self._extract_pages(storage_path, content_type)
            
            # 2. Chunk text and prepare payloads
            chunks_to_insert = []
            chunk_texts = []
            
            for page_num, page_text in pages:
                split_texts = self.splitter.split_text(page_text)
                for text_slice in split_texts:
                    chunk_texts.append(text_slice)
                    chunks_to_insert.append({
                        "chunk_id": uuid.uuid4(),
                        "content": text_slice,
                        "page_number": page_num,
                        "filename": filename
                    })
            
            if not chunks_to_insert:
                raise ValueError("No text could be extracted from this document.")
                
            # 3. Generate Embeddings (batch call)
            vectors = await self.embedding_service.get_embeddings(chunk_texts)
            for i, vector in enumerate(vectors):
                chunks_to_insert[i]["vector"] = vector
                
            # 4. Save chunks to Qdrant
            await self.vector_repo.upsert_chunks(doc_id, chunks_to_insert)
            
            # 5. Mark Document as COMPLETED
            await self.doc_repo.update_status(doc_id, DocumentStatus.COMPLETED)
            logger.info(f"Successfully processed document: {filename} (ID: {doc_id})")
            
        except Exception as e:
            logger.error(f"Failed to process document {filename}: {e}")
            # Mark Document as FAILED
            await self.doc_repo.update_status(
                doc_id, DocumentStatus.FAILED, error_message=str(e)
            )


    def _extract_pages(self, file_path: str, content_type: str) -> List[Tuple[Optional[int], str]]:
        pages = []
        
        # 1. PDF Parser
        if content_type == "application/pdf" or file_path.lower().endswith(".pdf"):
            try:
                from pypdf import PdfReader
                reader = PdfReader(file_path)
                for i, page in enumerate(reader.pages):
                    text = page.extract_text()
                    if text and text.strip():
                        pages.append((i + 1, text))
            except ImportError:
                logger.error("pypdf is not installed but a PDF file was uploaded.")
                raise ValueError("PDF parser not configured on server.")
                
        # 2. DOCX Parser
        elif (
            content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            or file_path.lower().endswith(".docx")
        ):
            try:
                import docx
                doc = docx.Document(file_path)
                text_list = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        text_list.append(para.text)
                
                for table in doc.tables:
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_text:
                            text_list.append(" | ".join(row_text))
                            
                full_text = "\n".join(text_list)
                if full_text.strip():
                    pages.append((None, full_text))
            except Exception as e:
                logger.error(f"Failed to parse DOCX file: {e}")
                raise ValueError(f"DOCX parser failed: {str(e)}")
                
        # 3. Text / Markdown Parser
        else:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
                if content.strip():
                    # Text files have no page numbers, set to None
                    pages.append((None, content))
                    
        return pages
